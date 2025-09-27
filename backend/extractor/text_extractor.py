"""
Text Extractor for Smart Doc Checker
Handles extraction of plain text from PDF and Word documents
"""

import os
import logging
from pathlib import Path
from typing import Dict, Optional

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


class DocumentExtractor:
    """Extracts text from various document formats"""
    
    def __init__(self):
        """Initialize the document extractor"""
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        
        # Check which libraries are available
        self.has_pdfplumber = pdfplumber is not None
        self.has_pymupdf = fitz is not None
        self.has_docx = Document is not None
        
        if not any([self.has_pdfplumber, self.has_pymupdf]):
            logging.warning("No PDF extraction libraries available. Install pdfplumber or PyMuPDF.")
        
        if not self.has_docx:
            logging.warning("python-docx not available. Cannot extract from Word documents.")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted plain text as string
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
            Exception: For other extraction errors
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_word(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"No extraction method for {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {file_path}: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using available libraries"""
        
        # Try pdfplumber first (generally better for structured documents)
        if self.has_pdfplumber:
            try:
                return self._extract_pdf_pdfplumber(file_path)
            except Exception as e:
                logging.warning(f"pdfplumber failed for {file_path}: {e}")
        
        # Fallback to PyMuPDF
        if self.has_pymupdf:
            try:
                return self._extract_pdf_pymupdf(file_path)
            except Exception as e:
                logging.warning(f"PyMuPDF failed for {file_path}: {e}")
        
        raise Exception("No PDF extraction libraries available or all failed")
    
    def _extract_pdf_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber"""
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}\n")
                except Exception as e:
                    logging.warning(f"Failed to extract page {page_num}: {e}")
                    continue
        
        if not text_content:
            raise Exception("No text could be extracted from PDF")
        
        return "\n".join(text_content)
    
    def _extract_pdf_pymupdf(self, file_path: str) -> str:
        """Extract text using PyMuPDF"""
        text_content = []
        
        pdf_document = fitz.open(file_path)
        
        try:
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}\n")
        finally:
            pdf_document.close()
        
        if not text_content:
            raise Exception("No text could be extracted from PDF")
        
        return "\n".join(text_content)
    
    def _extract_from_word(self, file_path: str) -> str:
        """Extract text from Word document"""
        
        if not self.has_docx:
            raise Exception("python-docx library not available")
        
        try:
            # Handle both .docx and .doc (though .doc support is limited)
            doc = Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise Exception("No text could be extracted from Word document")
            
            return "\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract from Word document: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise Exception(f"Failed to read text file with multiple encodings: {str(e)}")
    
    def get_document_info(self, file_path: str) -> Dict[str, any]:
        """
        Get basic information about a document
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with document information
        """
        if not os.path.exists(file_path):
            return {"error": "File not found"}
        
        file_stats = os.stat(file_path)
        file_path_obj = Path(file_path)
        
        info = {
            "filename": file_path_obj.name,
            "file_size": file_stats.st_size,
            "file_extension": file_path_obj.suffix.lower(),
            "is_supported": file_path_obj.suffix.lower() in self.supported_formats,
            "last_modified": file_stats.st_mtime
        }
        
        # Add format-specific information
        if file_path_obj.suffix.lower() == '.pdf' and self.has_pymupdf:
            try:
                pdf_doc = fitz.open(file_path)
                info["page_count"] = pdf_doc.page_count
                info["pdf_info"] = pdf_doc.metadata
                pdf_doc.close()
            except:
                pass
        
        elif file_path_obj.suffix.lower() in ['.docx'] and self.has_docx:
            try:
                doc = Document(file_path)
                info["paragraph_count"] = len(doc.paragraphs)
                info["table_count"] = len(doc.tables)
            except:
                pass
        
        return info
    
    def batch_extract(self, file_paths: list) -> Dict[str, any]:
        """
        Extract text from multiple documents
        
        Args:
            file_paths: List of file paths
            
        Returns:
            Dictionary with results for each file
        """
        results = {
            "successful": [],
            "failed": [],
            "summary": {
                "total_files": len(file_paths),
                "successful_count": 0,
                "failed_count": 0
            }
        }
        
        for file_path in file_paths:
            try:
                text = self.extract_text(file_path)
                results["successful"].append({
                    "file_path": file_path,
                    "filename": os.path.basename(file_path),
                    "text": text,
                    "text_length": len(text),
                    "status": "success"
                })
                results["summary"]["successful_count"] += 1
                
            except Exception as e:
                results["failed"].append({
                    "file_path": file_path,
                    "filename": os.path.basename(file_path),
                    "error": str(e),
                    "status": "failed"
                })
                results["summary"]["failed_count"] += 1
        
        return results


# Test the extractor
if __name__ == "__main__":
    print("📄 Testing Document Extractor")
    print("=" * 35)
    
    extractor = DocumentExtractor()
    
    print(f"Supported formats: {extractor.supported_formats}")
    print(f"Available libraries:")
    print(f"  - pdfplumber: {extractor.has_pdfplumber}")
    print(f"  - PyMuPDF: {extractor.has_pymupdf}")
    print(f"  - python-docx: {extractor.has_docx}")
    
    # Test with sample files if they exist
    uploads_dir = Path(__file__).parent.parent.parent / "uploads"
    if uploads_dir.exists():
        sample_files = list(uploads_dir.glob("*.*"))
        supported_files = [f for f in sample_files if f.suffix.lower() in extractor.supported_formats]
        
        if supported_files:
            print(f"\nTesting with {len(supported_files)} files:")
            results = extractor.batch_extract([str(f) for f in supported_files])
            
            print(f"✅ Successful: {results['summary']['successful_count']}")
            print(f"❌ Failed: {results['summary']['failed_count']}")
            
            for result in results["successful"]:
                print(f"  - {result['filename']}: {result['text_length']} characters")
                
            for result in results["failed"]:
                print(f"  - {result['filename']}: {result['error']}")
        else:
            print("\nNo supported files found in uploads directory")
    else:
        print("\nUploads directory not found")
    
    print("\n✅ Extractor test completed!")